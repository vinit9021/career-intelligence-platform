from io import BytesIO
from typing import Any
from uuid import uuid4

import docx
import pytest
from httpx2 import AsyncClient


def build_resume_docx() -> bytes:
    document = docx.Document()
    document.add_paragraph("Professional Summary")
    document.add_paragraph("Backend engineer building reliable APIs")
    document.add_paragraph("Technical Skills")
    document.add_paragraph("Python, FastAPI, PostgreSQL")
    document.add_paragraph("Education")
    document.add_paragraph("B.Tech - IIT Ropar")
    document.add_paragraph("Experience")
    document.add_paragraph("Software Engineering Intern")
    document.add_paragraph("Projects")
    document.add_paragraph("Career Intelligence Platform")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def register_user(client: AsyncClient) -> dict[str, Any]:
    response = await client.post(
        "/api/v1/auth/register",
        json={
            "email": f"phase1-{uuid4()}@example.com",
            "password": "StrongPassword9!",
            "full_name": "Phase One User",
        },
    )

    assert response.status_code == 201
    data: dict[str, Any] = response.json()
    return data


def auth_headers(registration: dict[str, Any]) -> dict[str, str]:
    return {"Authorization": f"Bearer {registration['access_token']}"}


@pytest.mark.asyncio
async def test_complete_phase1_resume_lifecycle(
    auth_client: AsyncClient,
) -> None:
    registration = await register_user(auth_client)
    headers = auth_headers(registration)
    original_data = build_resume_docx()

    upload = await auth_client.post(
        "/api/v1/resume/upload",
        headers=headers,
        files={
            "file": (
                "phase-one-résumé.docx",
                original_data,
                ("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            )
        },
    )
    assert upload.status_code == 201
    resume_id = str(upload.json()["id"])

    pending = await auth_client.get(
        f"/api/v1/resume/{resume_id}/parse-status",
        headers=headers,
    )
    assert pending.status_code == 200
    assert pending.json()["status"] == "pending"

    parsed = await auth_client.post(
        f"/api/v1/resume/{resume_id}/parse",
        headers=headers,
    )
    assert parsed.status_code == 200
    assert parsed.json()["status"] == "completed"

    viewer = await auth_client.get(
        f"/api/v1/resume/{resume_id}/viewer",
        headers=headers,
    )
    assert viewer.status_code == 200
    assert viewer.json()["content"]["skills"] == [
        "Python",
        "FastAPI",
        "PostgreSQL",
    ]

    history = await auth_client.get(
        "/api/v1/resume/history?page=1&page_size=20",
        headers=headers,
    )
    assert history.status_code == 200
    assert history.json()["total"] == 1
    assert history.json()["items"][0]["id"] == resume_id

    file_response = await auth_client.get(
        f"/api/v1/resume/{resume_id}/file",
        headers=headers,
    )
    assert file_response.status_code == 200
    assert file_response.content == original_data
    assert file_response.headers["cache-control"] == "private, no-store"
    assert file_response.headers["pragma"] == "no-cache"
    assert file_response.headers["x-content-type-options"] == "nosniff"
    assert file_response.headers["content-length"] == str(len(original_data))
    assert len(file_response.headers["x-content-sha256"]) == 64
    assert (
        "filename*=UTF-8''phase-one-r%C3%A9sum%C3%A9.docx"
        in (file_response.headers["content-disposition"])
    )

    deleted = await auth_client.delete(
        f"/api/v1/resume/{resume_id}",
        headers=headers,
    )
    assert deleted.status_code == 200
    assert deleted.json()["deleted"] is True

    missing = await auth_client.get(
        f"/api/v1/resume/{resume_id}",
        headers=headers,
    )
    assert missing.status_code == 404

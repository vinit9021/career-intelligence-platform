from io import BytesIO
from pathlib import Path
from uuid import uuid4

import docx
import pytest
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker

from app.models import Resume, User
from app.parsers import build_default_parser_registry
from app.services.resume_parsing import (
    ResumeNotFoundError,
    ResumeParseResultNotFoundError,
    ResumeParsingService,
    ResumeSourceUnavailableError,
)
from app.storage import LocalStorage

SessionFactory = async_sessionmaker[AsyncSession]


def build_docx() -> bytes:
    document = docx.Document()
    document.add_paragraph("Professional Summary")
    document.add_paragraph("Backend engineer")
    document.add_paragraph("Technical Skills")
    document.add_paragraph("Python, FastAPI, PostgreSQL")
    document.add_paragraph("Education")
    document.add_paragraph("B.Tech - IIT Ropar")
    document.add_paragraph("Experience")
    document.add_paragraph("Software Engineering Intern")
    document.add_paragraph("Projects")
    document.add_paragraph("Career Intelligence Platform")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def create_resume(
    session: AsyncSession,
    storage: LocalStorage,
) -> tuple[User, Resume]:
    user = User(
        id=uuid4(),
        email="alice@example.com",
        password_hash="stored-password-hash",
        full_name="Alice Example",
    )
    resume_id = uuid4()
    key = f"resumes/{user.id}/{resume_id}/original.docx"
    data = build_docx()
    await storage.save(
        key=key,
        data=data,
        content_type=("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        checksum_sha256="a" * 64,
    )
    resume = Resume(
        id=resume_id,
        user_id=user.id,
        original_filename="resume.docx",
        storage_backend="local",
        storage_key=key,
        storage_etag="a" * 64,
        content_type=("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        file_extension="docx",
        file_size_bytes=len(data),
        sha256="a" * 64,
    )
    session.add_all([user, resume])
    await session.commit()
    return user, resume


@pytest.mark.asyncio
async def test_parse_resume_persists_structured_result(
    test_session_factory: SessionFactory,
    tmp_path: Path,
) -> None:
    storage = LocalStorage(tmp_path / "storage")

    async with test_session_factory() as session:
        user, resume = await create_resume(session, storage)
        service = ResumeParsingService(
            session=session,
            storage=storage,
            parser_registry=build_default_parser_registry(),
        )

        response = await service.parse_resume(user=user, resume_id=resume.id)
        stored = await service.get_parse_result(user=user, resume_id=resume.id)

        assert response.status == "completed"
        assert response.content.skills == ["Python", "FastAPI", "PostgreSQL"]
        assert stored.id == response.id
        assert stored.raw_text.startswith("Professional Summary")
        assert resume.parse_status == "completed"
        assert resume.parsed_at is not None


@pytest.mark.asyncio
async def test_missing_resume_is_rejected(
    test_session_factory: SessionFactory,
    tmp_path: Path,
) -> None:
    storage = LocalStorage(tmp_path / "storage")

    async with test_session_factory() as session:
        user = User(
            id=uuid4(),
            email="alice@example.com",
            password_hash="stored-password-hash",
            full_name="Alice Example",
        )
        session.add(user)
        await session.commit()
        service = ResumeParsingService(
            session=session,
            storage=storage,
            parser_registry=build_default_parser_registry(),
        )

        with pytest.raises(ResumeNotFoundError):
            await service.parse_resume(user=user, resume_id=uuid4())


@pytest.mark.asyncio
async def test_missing_parse_result_is_rejected(
    test_session_factory: SessionFactory,
    tmp_path: Path,
) -> None:
    storage = LocalStorage(tmp_path / "storage")

    async with test_session_factory() as session:
        user, resume = await create_resume(session, storage)
        service = ResumeParsingService(
            session=session,
            storage=storage,
            parser_registry=build_default_parser_registry(),
        )

        with pytest.raises(ResumeParseResultNotFoundError):
            await service.get_parse_result(user=user, resume_id=resume.id)


@pytest.mark.asyncio
async def test_missing_stored_file_marks_resume_failed(
    test_session_factory: SessionFactory,
    tmp_path: Path,
) -> None:
    storage = LocalStorage(tmp_path / "storage")

    async with test_session_factory() as session:
        user, resume = await create_resume(session, storage)
        await storage.delete(key=resume.storage_key)
        service = ResumeParsingService(
            session=session,
            storage=storage,
            parser_registry=build_default_parser_registry(),
        )

        with pytest.raises(ResumeSourceUnavailableError):
            await service.parse_resume(user=user, resume_id=resume.id)

        assert resume.parse_status == "failed"
        assert resume.parse_error == "The stored resume could not be read."

from io import BytesIO
from typing import Any

import docx
import pytest
from httpx2 import AsyncClient


def build_docx() -> bytes:
    document = docx.Document()
    document.add_paragraph("Professional Summary")
    document.add_paragraph("Backend engineer")
    document.add_paragraph("Technical Skills")
    document.add_paragraph("Python, FastAPI, PostgreSQL")
    document.add_paragraph("Education")
    document.add_paragraph("B.Tech - IIT Ropar")
    document.add_paragraph("Experience")
    document.add_paragraph("Software Engineering Intern")
    document.add_paragraph("Projects")
    document.add_paragraph("Career Intelligence Platform")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def register_user(client: AsyncClient) -> dict[str, Any]:
    response = await client.post(
        "/api/v1/auth/register",
        json={
            "email": "alice@example.com",
            "password": "StrongPassword9!",
            "full_name": "Alice Example",
        },
    )
    assert response.status_code == 201
    data: dict[str, Any] = response.json()
    return data


def auth_headers(registration: dict[str, Any]) -> dict[str, str]:
    return {"Authorization": f"Bearer {registration['access_token']}"}


@pytest.mark.asyncio
async def test_authenticated_user_can_parse_uploaded_resume(
    auth_client: AsyncClient,
) -> None:
    registration = await register_user(auth_client)
    headers = auth_headers(registration)
    upload = await auth_client.post(
        "/api/v1/resume/upload",
        headers=headers,
        files={
            "file": (
                "resume.docx",
                build_docx(),
                ("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            )
        },
    )
    assert upload.status_code == 201
    resume_id = upload.json()["id"]

    parsed = await auth_client.post(
        f"/api/v1/resume/{resume_id}/parse",
        headers=headers,
    )
    assert parsed.status_code == 200
    body = parsed.json()
    assert body["status"] == "completed"
    assert body["content"]["skills"] == ["Python", "FastAPI", "PostgreSQL"]

    fetched = await auth_client.get(
        f"/api/v1/resume/{resume_id}/parsed",
        headers=headers,
    )
    assert fetched.status_code == 200
    assert fetched.json()["id"] == body["id"]


@pytest.mark.asyncio
async def test_parse_endpoint_requires_authentication(
    auth_client: AsyncClient,
) -> None:
    response = await auth_client.post("/api/v1/resume/00000000-0000-0000-0000-000000000000/parse")

    assert response.status_code == 401

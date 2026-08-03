from io import BytesIO

import docx
import pytest

from app.parsers import ResumeParserError
from app.parsers.docx import DocxResumeParser


def build_docx() -> bytes:
    document = docx.Document()
    document.add_paragraph("Summary")
    document.add_paragraph("Backend engineer")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "FastAPI"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_docx_parser_preserves_paragraph_and_table_order() -> None:
    result = DocxResumeParser().parse(build_docx())

    assert result.page_count is None
    assert result.raw_text.splitlines() == [
        "Summary",
        "Backend engineer",
        "Python | FastAPI",
    ]
    assert result.requires_ocr is False


def test_invalid_docx_is_rejected() -> None:
    with pytest.raises(ResumeParserError):
        DocxResumeParser().parse(b"not-a-docx")

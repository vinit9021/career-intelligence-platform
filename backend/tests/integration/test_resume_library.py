from io import BytesIO
from typing import Any

import docx
import pytest
from httpx2 import AsyncClient


def build_docx() -> bytes:
    document = docx.Document()
    document.add_paragraph("Professional Summary")
    document.add_paragraph("Backend engineer")
    document.add_paragraph("Technical Skills")
    document.add_paragraph("Python, FastAPI, PostgreSQL")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def register_user(
    client: AsyncClient,
    *,
    email: str,
) -> dict[str, Any]:
    response = await client.post(
        "/api/v1/auth/register",
        json={
            "email": email,
            "password": "StrongPassword9!",
            "full_name": "Resume User",
        },
    )
    assert response.status_code == 201
    data: dict[str, Any] = response.json()
    return data


def auth_headers(registration: dict[str, Any]) -> dict[str, str]:
    return {"Authorization": f"Bearer {registration['access_token']}"}


async def upload_resume(
    client: AsyncClient,
    *,
    headers: dict[str, str],
    filename: str = "resume.docx",
) -> tuple[str, bytes]:
    data = build_docx()
    response = await client.post(
        "/api/v1/resume/upload",
        headers=headers,
        files={
            "file": (
                filename,
                data,
                ("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            )
        },
    )
    assert response.status_code == 201
    return str(response.json()["id"]), data


@pytest.mark.asyncio
async def test_history_detail_and_pending_status(
    auth_client: AsyncClient,
) -> None:
    registration = await register_user(
        auth_client,
        email="history@example.com",
    )
    headers = auth_headers(registration)
    first_id, _ = await upload_resume(
        auth_client,
        headers=headers,
        filename="first.docx",
    )
    second_id, _ = await upload_resume(
        auth_client,
        headers=headers,
        filename="second.docx",
    )

    history = await auth_client.get(
        "/api/v1/resume/history?page=1&page_size=1",
        headers=headers,
    )
    assert history.status_code == 200
    history_body = history.json()
    assert history_body["total"] == 2
    assert history_body["total_pages"] == 2
    assert len(history_body["items"]) == 1

    detail = await auth_client.get(
        f"/api/v1/resume/{first_id}",
        headers=headers,
    )
    assert detail.status_code == 200
    assert detail.json()["id"] == first_id
    assert "storage_key" not in detail.json()

    parse_status = await auth_client.get(
        f"/api/v1/resume/{second_id}/parse-status",
        headers=headers,
    )
    assert parse_status.status_code == 200
    assert parse_status.json()["status"] == "pending"
    assert parse_status.json()["has_parsed_result"] is False


@pytest.mark.asyncio
async def test_viewer_file_and_delete_lifecycle(
    auth_client: AsyncClient,
) -> None:
    registration = await register_user(
        auth_client,
        email="viewer@example.com",
    )
    headers = auth_headers(registration)
    resume_id, original_data = await upload_resume(
        auth_client,
        headers=headers,
    )

    parsed = await auth_client.post(
        f"/api/v1/resume/{resume_id}/parse",
        headers=headers,
    )
    assert parsed.status_code == 200

    viewer = await auth_client.get(
        f"/api/v1/resume/{resume_id}/viewer",
        headers=headers,
    )
    assert viewer.status_code == 200
    assert viewer.json()["content"]["skills"] == [
        "Python",
        "FastAPI",
        "PostgreSQL",
    ]

    file_response = await auth_client.get(
        f"/api/v1/resume/{resume_id}/file",
        headers=headers,
    )
    assert file_response.status_code == 200
    assert file_response.content == original_data
    assert "filename*=UTF-8''resume.docx" in file_response.headers["content-disposition"]
    assert len(file_response.headers["x-content-sha256"]) == 64

    deleted = await auth_client.delete(
        f"/api/v1/resume/{resume_id}",
        headers=headers,
    )
    assert deleted.status_code == 200
    assert deleted.json()["deleted"] is True

    missing = await auth_client.get(
        f"/api/v1/resume/{resume_id}",
        headers=headers,
    )
    assert missing.status_code == 404


@pytest.mark.asyncio
async def test_resume_ownership_is_enforced(
    auth_client: AsyncClient,
) -> None:
    owner = await register_user(
        auth_client,
        email="owner@example.com",
    )
    other = await register_user(
        auth_client,
        email="other@example.com",
    )
    owner_headers = auth_headers(owner)
    other_headers = auth_headers(other)
    resume_id, _ = await upload_resume(
        auth_client,
        headers=owner_headers,
    )

    for method, path in (
        ("GET", f"/api/v1/resume/{resume_id}"),
        ("GET", f"/api/v1/resume/{resume_id}/parse-status"),
        ("GET", f"/api/v1/resume/{resume_id}/viewer"),
        ("GET", f"/api/v1/resume/{resume_id}/file"),
        ("DELETE", f"/api/v1/resume/{resume_id}"),
    ):
        response = await auth_client.request(
            method,
            path,
            headers=other_headers,
        )
        assert response.status_code == 404


@pytest.mark.asyncio
async def test_history_requires_authentication_and_valid_pagination(
    auth_client: AsyncClient,
) -> None:
    unauthorized = await auth_client.get("/api/v1/resume/history")
    assert unauthorized.status_code == 401

    registration = await register_user(
        auth_client,
        email="pagination@example.com",
    )
    invalid = await auth_client.get(
        "/api/v1/resume/history?page=0&page_size=101",
        headers=auth_headers(registration),
    )
    assert invalid.status_code == 422
